"""DOCX text and image extraction with vision processing."""

from typing import Dict, Any
from docx import Document


def extract_docx(filepath: str, process_images: bool = True) -> Dict[str, Any]:
    """
    Extract content from Word document.

    Args:
        filepath: Path to .docx file
        process_images: If True, process images with vision

    Returns:
        Dict with content, images, metadata
    """
    if process_images:
        from ..services.vision import process_image

    doc = Document(filepath)

    # Extract text from paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    content = "\n\n".join(paragraphs)

    # Extract images
    images = []

    if process_images and hasattr(doc, 'inline_shapes'):
        for img_index, inline_shape in enumerate(doc.inline_shapes):
            try:
                if inline_shape.type == 3:  # PICTURE type
                    image_bytes = inline_shape.image.blob
                    vision_result = process_image(image_bytes)

                    images.append({
                        "image_index": img_index,
                        "type": vision_result["type"],
                        "quality_score": vision_result["quality_score"],
                        "extracted_text": vision_result["extracted_text"],
                        "should_index": vision_result["should_index"]
                    })
            except Exception as e:
                print(f"Warning: Failed to process inline shape {img_index}: {e}")

    return {
        "content": content,
        "images": images,
        "metadata": {
            "total_paragraphs": len(doc.paragraphs),
            "images_processed": len(images)
        }
    }
